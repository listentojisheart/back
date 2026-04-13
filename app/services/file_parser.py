"""
File parsing: extract text from PDF, DOCX, TXT, MD, YAML.
"""
import io
from pathlib import Path


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """Return extracted text from a file blob. Best-effort across formats."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(content)
    if ext == ".docx":
        return _extract_docx(content)
    if ext in (".txt", ".md", ".yaml", ".yml"):
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            return content.decode("latin-1", errors="replace")
    # fallback: try utf-8
    try:
        return content.decode("utf-8", errors="replace")
    except Exception:
        return ""


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        pages_text = []
        for page in reader.pages:
            try:
                pages_text.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n\n".join(pages_text)
    except Exception as e:
        return f"[PDF extraction failed: {e}]"


def _extract_docx(content: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        paras = [p.text for p in doc.paragraphs if p.text]
        # Also pull table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paras.append(cell.text)
        return "\n".join(paras)
    except Exception as e:
        return f"[DOCX extraction failed: {e}]"
